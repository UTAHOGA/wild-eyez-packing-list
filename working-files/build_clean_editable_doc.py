from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\tyler\GitHub\wild-eyez-packing-list")
WORK = ROOT / "working-files"
ASSETS = ROOT / "assets"
OUT_DOCX = WORK / "wild_eyez_packing_list_CLEAN_EDITABLE_SOURCE.docx"
AUDIO_URL = "https://drive.google.com/file/d/109ytIBbHBxUSXiNJerwfcgbhrrcFbxu9/view?usp=sharing"

GREEN = "173C2C"
TAN = "E8DCC8"
BRONZE = "B48A54"
MAROON = "7D2530"
PINK = "F5DFE3"


def set_font(run, size=None, bold=None, italic=None, color=None, name="Georgia"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for edge, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def cell_border(cell, top=None, left=None, bottom=None, right=None):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge, data in [("top", top), ("left", left), ("bottom", bottom), ("right", right)]:
        if not data:
            continue
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        for key, value in data.items():
            node.set(qn(f"w:{key}"), str(value))


def table_borders(table, color="E5DAC8", sz="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), sz)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def paragraph(container, text="", size=10.1, bold=False, italic=False, color="222222",
              align=None, before=0, after=5, font="Georgia"):
    p = container.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.05
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_font(r, size=size, bold=bold, italic=italic, color=color, name=font)
    return p


def heading(doc, text, size=16):
    p = paragraph(doc, "", before=6, after=5)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_font(r, size=size, bold=True, color=GREEN, name="Arial")


def banner(doc, label):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(7.35)
    cell = table.cell(0, 0)
    shade(cell, GREEN)
    cell_margins(cell, 115, 180, 115, 180)
    cell_border(
        cell,
        top={"val": "single", "sz": "8", "color": BRONZE},
        left={"val": "single", "sz": "8", "color": GREEN},
        bottom={"val": "single", "sz": "8", "color": BRONZE},
        right={"val": "single", "sz": "8", "color": GREEN},
    )
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("WILD EYEZ OUTFITTERS")
    set_font(r, size=14, bold=True, italic=True, color="F4E6CC", name="Georgia")
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Client Western Big Game Packing List & Preparation Guide")
    set_font(r2, size=10, bold=True, color="FFFFFF", name="Arial")
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(label)
    set_font(r3, size=8.5, italic=True, color=TAN, name="Arial")
    paragraph(doc, "", after=2)


def callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(7.15)
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    shade(cell, PINK)
    cell_margins(cell, 135, 185, 135, 185)
    cell_border(
        cell,
        top={"val": "single", "sz": "6", "color": MAROON},
        left={"val": "single", "sz": "28", "color": MAROON},
        bottom={"val": "single", "sz": "6", "color": MAROON},
        right={"val": "single", "sz": "6", "color": MAROON},
    )
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title.upper())
    set_font(r, size=12, bold=True, color=MAROON, name="Arial")
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.03
    r2 = p2.add_run(body)
    set_font(r2, size=10, color="111111", name="Georgia")
    paragraph(doc, "", after=2)


def checklist(doc, title, items):
    p = paragraph(doc, "", before=3, after=2)
    r = p.add_run(title.upper())
    set_font(r, size=11.5, bold=True, color=GREEN, name="Arial")
    table = doc.add_table(rows=len(items), cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(7.15)
    table_borders(table)
    for idx, item in enumerate(items):
        cell = table.cell(idx, 0)
        shade(cell, "FBF8F1" if idx % 2 == 0 else "F2EBDD")
        cell_margins(cell, 78, 100, 78, 100)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run("[ ]  ")
        set_font(r, size=10.7, bold=True, color="111111", name="Arial")
        r2 = p.add_run(item)
        set_font(r2, size=9.5, color="111111", name="Arial")


def add_link(paragraph_obj, text, url):
    r_id = paragraph_obj.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), MAROON)
    r_pr.append(color)
    bold = OxmlElement("w:b")
    r_pr.append(bold)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "none")
    r_pr.append(underline)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph_obj._p.append(hyperlink)


def page_break(doc):
    doc.add_page_break()


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.45)
section.bottom_margin = Inches(0.45)
section.left_margin = Inches(0.55)
section.right_margin = Inches(0.55)

doc.styles["Normal"].font.name = "Georgia"
doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Georgia")
doc.styles["Normal"].font.size = Pt(10.1)

banner(doc, "PREPARATION | KNOWLEDGE | EXPERIENCE | SAFETY + ETHICS = SUCCESS")
logo = ASSETS / "wild-eyez-logo.png"
if logo.exists():
    p = paragraph(doc, "", after=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(logo), width=Inches(2.15))
heading(doc, "CLIENT WESTERN BIG GAME PACKING LIST & PREPARATION GUIDE", 16)
p = paragraph(doc, "", after=4)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("~SEE THE WILD!!!")
set_font(r, size=14, bold=True, italic=True, color=MAROON, name="Georgia")

logo_table = doc.add_table(rows=1, cols=4)
logo_table.alignment = WD_TABLE_ALIGNMENT.CENTER
logo_table.autofit = False
for idx, image in enumerate([ASSETS / "usfs.png", ASSETS / "blm.png", ASSETS / "dwr.png", ASSETS / "uoga.png"]):
    logo_table.columns[idx].width = Inches(1.1)
    cell = logo_table.cell(0, idx)
    shade(cell, "FFFFFF")
    cell_margins(cell, 35, 35, 35, 35)
    cell_border(
        cell,
        top={"val": "single", "sz": "3", "color": "CDB98F"},
        left={"val": "single", "sz": "3", "color": "CDB98F"},
        bottom={"val": "single", "sz": "3", "color": "CDB98F"},
        right={"val": "single", "sz": "3", "color": "CDB98F"},
    )
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if image.exists():
        p.add_run().add_picture(str(image), width=Inches(0.72))

heading(doc, "Welcome", 13)
paragraph(doc, "I'm Tyler Miller with Wild Eyez Outfitters. This Client Western Big Game Packing List and Preparation Guide provides a general packing list, preparation guide, and outline of what to expect during a big game hunt in Utah. It includes important safety information, equipment recommendations, clothing suggestions, and practical guidance for preparing for your hunt. If you received this guide, you may be considering a hunt with us, or you may have already contracted your hunt. Either way, I want to thank you for considering Wild Eyez Outfitters and congratulate you on your upcoming hunting adventure.")
paragraph(doc, "While we keep tabs on your target big game animal, this is a short outline of what we expect of you and expect you to show up with in preparation for this rare opportunity that you are fortunate enough to have this fall.")
paragraph(doc, "The items in this guide are not brand specific. I may occasionally recommend a particular brand or explain whether an item is essential or optional. However, the final selection of clothing and equipment remains at your discretion.")
paragraph(doc, "Utah's weather, terrain, elevation, and big game habitat vary greatly. The information in this guide is based on knowledge of the local area, seasonal migration patterns, previous client hunts, and more than forty years of personal hunting and field experience.")
p = paragraph(doc, "", after=6)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_link(p, "AUDIO REFERENCE", AUDIO_URL)

page_break(doc)
banner(doc, "SAFETY FIRST")
callout(doc, "Priority #1 - Your Safety", "Conditions may become extreme in our hunting environment. Know your own personal limitations and follow all health practitioner advice that has been given to you above what your hunting guide may suggest or expect of you. Big game hunting is an extreme, adventurous sport. It is demanding both physically and mentally. Pursuing and taking of a big game animal is an intense and rigorous activity.")
paragraph(doc, "Your Guide(s) and I will watch you closely, but we are VERY focused on the pursuit of your game animal, so it is your responsibility to know and watch yourself. We cannot see or feel the world through your eyes, and we operate at or near our own personal limits often when pursuing big game. We cannot be liable if you choose to exceed your own personal abilities and limitations. Reality is that you need to prepare appropriately or your opportunity for success will suffer.")
paragraph(doc, "Hunting is LETHAL by nature; something will likely die. We want that to be the game animal that is on your permit and not yourself or any other unintended life.", bold=True)
heading(doc, "What To Expect", 13)
paragraph(doc, "Expect travel methods to vary with weather, terrain, and access conditions. We may use a heated 4x4, snowmobiles, ATVs or side-by-sides, equine travel, or hiking, but our primary way of reaching hunting ground is usually on foot. If you are uncomfortable with any method of travel, please let me know before your hunt so I can plan accordingly.")
checklist(doc, "Possible Travel Methods", [
    "A heated 4x4 off-road hunting automobile, if conditions permit and it is the best choice for access",
    "Snowmobiles when snow and weather require them",
    "ATV's (4-wheelers) and side by side OHV where terrain and conditions allow",
    "Equine animal when that is the most practical travel method for the area",
    "Hiking, which is still the primary way we get where we need to go",
])
paragraph(doc, "A typical day will start around 5:00 a.m. We will develop a plan based on current big game animal location, weather conditions, hunting conditions, and past experience. A typical pursuit may include traveling to the intended hunting location, spotting and glassing for the game animal, and then executing a stalk. This stalk may involve hiking, sitting, kneeling, spotting, and moving again while carrying your weapon, gear, and pack.")
callout(doc, "Professional Pointer", "When conditions allow, plan to sleep or fully rest during the middle of the day. If midday hunting is productive, we may remain active, continue glassing, reposition, or continue the pursuit. When downtime is available, use it to reset your mind. The goal is to remain mentally sharp when the evening opportunity arrives.")

page_break(doc)
banner(doc, "WEATHER | TERRAIN | OPPORTUNITY")
heading(doc, "Weather & Terrain", 13)
paragraph(doc, "We may hunt pinon and juniper country at roughly 5,000 to 7,000 feet, or quaking aspen and blue spruce terrain from roughly 7,000 to 11,000 feet. Cover may range from semi-open to dense. Distance to your game can vary from 50 to 800 yards or more. You should be comfortable lying down, sitting, kneeling, and standing while establishing a secure rest. The prone position is usually the most stable.")
paragraph(doc, "One suggestion is to carry a waterproof mat or similar pad which can be very quickly deployed to lay upon rather than directly in the dirt, mud or snow. However, it is not always possible to deploy this equipment before taking your shot and you should consider it an exception to the general rule. Plan on and prepare for weather conditions from hot afternoons on rocky, dry dirt terrain down to extremes of snow conditions and freezing temperatures. A hunt may move from 70- or 80-degree dry terrain to a foot of snow and freezing conditions. Weather can remain consistent or change dramatically in less than one hour. That variability is manageable when you prepare for it.")
callout(doc, "Be Ready When The Opportunity Changes", "A careful stalk may give us time to study the animal, build a rest, and set up video equipment. Just as often, the opportunity develops in seconds. You must be able to deploy your rifle, establish a stable rest, and prepare for the shot without learning your equipment in the moment. I will provide the range, but you must know your rifle well enough to make the proper distance and wind adjustment.")
heading(doc, "Clothing & Footwear", 13)
paragraph(doc, "Do not skimp on clothing. You may have only seconds to sit or lie in snow, gain your sight picture, and take the shot. Your outer layer should be waterproof or highly water resistant. Build a quiet, flexible layering system that allows you to add, remove, or vent clothing as conditions change. Avoid cotton, which holds moisture, and test every garment by rubbing the fabric together. If it whistles, crackles, or rubs loudly, it can compromise a stalk. Overdress rather than underdress; you can remove a layer, but you cannot put on clothing you left at home.")
callout(doc, "Professional Pointer", "Rub the arms, legs, and fabric panels of every garment together before packing it. If it makes excessive whistling, rubbing, plastic, ski-pant, or Levi-type noise, it is too loud for an effective stalk. Use quiet, mobile, flexible material.")

page_break(doc)
banner(doc, "CLOTHING & FOOTWEAR CHECKLIST")
checklist(doc, "Upper Body", [
    "Base layer: undershirt and thermal layer. These should be comfortable, quiet, and easy to move in.",
    "Middle layer: light and mid-weight shirt, light-weight vest, and a light jacket or hoodie.",
    "Top / outer layer: mid-weight jacket and heavy-weight coat with a hood.",
    "Additional waterproof or highly water-resistant outerwear for wet-weather conditions.",
])
checklist(doc, "Lower Body", [
    "Base layer: long underwear or thermal layer.",
    "Middle layer: joggers or light to mid-weight pants.",
    "Top / outer layer: mid-heavy weight pants with enough room to move comfortably.",
    "Additional lightweight or packable waterproof outerwear for wet-weather conditions.",
])
checklist(doc, "Feet, Head & Hands", [
    "Gaiters to keep dirt and snow out of your boots and off your lower legs.",
    "Two pairs of socks: one light weight and one merino wool, or one good heavy-weight merino wool pair if your boots are high quality.",
    "Waterproof boots with aggressive soles, already broken in before opening day.",
    "Backup snowmobile boots or insulated rubber boots with a waterproof, rubber-based sole when conditions call for them.",
    "Lightweight cap plus a heavyweight hunting cap, balaclava, or facemask.",
    "Four pairs of gloves: heavy waterproof gloves, medium working/hiking gloves, and two lightweight pairs you can fire your rifle with.",
    "Optional safety sunglasses, polarized with UV protection and light amber or brown transitional lenses.",
])
callout(doc, "Pro Tip - Protect Your Feet", "Protect your feet before the hunt starts. Break in your boots before opening day, test your sock system on real walks or hikes, and address hot spots immediately. A small blister on day one can become the reason you cannot hunt hard on day three. Your feet are part of your weapon system because they determine whether you can still get into position when the opportunity comes.")
paragraph(doc, 'Cotton clothing is highly discouraged. "Cotton Kills." For all clothing, wash it at least once in baking soda, not perfumed detergents, and allow it to air dry.', bold=True)

page_break(doc)
banner(doc, "FIELD GEAR")
paragraph(doc, "Every item in your pack should solve a real field problem: navigation, water, warmth, communication, meat care, repair, or sustained energy. Pack deliberately and know exactly where essential items are located. Practice retrieving them while wearing your hunting clothing. That may sound elementary, but when adrenaline hits, familiar movements matter. For hydration, I commonly use water mixed with an electrolyte drink, supported by high-carbohydrate foods the body can access quickly.")
checklist(doc, "Field Gear - Essentials", [
    "Hunting license, permit, identification, and any required hunter-safety documentation.",
    "Offline map, GPS or mapping app, and a backup compass or paper map.",
    "Satellite communicator, emergency locator, or other reliable backup communication option when cell service is limited.",
    "Headlamp plus extra batteries or a backup light. You may leave before daylight and return after dark.",
    "Compact first-aid and foot-care kit, including blister care, tape, pain reliever, and personal medications.",
    "Water bottle or bladder plus water treatment tablets, filter, or other backup water-purification method.",
    "Wind direction indicator, such as a wind puffer or lightweight indicator material.",
    "Binoculars, rangefinder, lens cloth, and any optics support you intend to use.",
    "Two sharp knives with a sharpening stone or similar sharpening device.",
    "Game bags. Preferably heavy canvas and not the netting type game bags. Potato or burlap sacks are cheap and perfect.",
    "Four pairs of nitrile or latex gloves. They are optional, but handy to have for cleanup after quartering your game.",
    "Backpack or daypack with a bladder. Do not go cheap on a backpack.",
    "Baby wipes.",
    "Waterproof mat or similar quick-deploy pad.",
    "Lighter and matches. Vaseline-saturated cotton balls work well for starting a fire in difficult situations.",
    "Small repair kit: duct tape or gear tape, extra boot laces, zip ties, and a compact multi-tool.",
    "High-carbohydrate food and fluids like candy bars, canned fruit, pudding, or sandwiches.",
])

page_break(doc)
banner(doc, "FIREARMS & AMMUNITION")
callout(doc, "Professional Pointer", "Walking sticks or shooting sticks are fine if desired, but they MUST BE wrapped in tape or similar silencing material on the ends. Walking sticks hitting a rock will send your trophy animal running the opposite direction before you even knew he was there.")
paragraph(doc, "Bring the rifle you know best and are most comfortable operating under pressure. Arrive with enough ammunition for the hunt and for meaningful practice before arrival. Your scope should be checked, your rifle should be zeroed before the hunt, and your system should include any bipod, independent shooting sticks, cleaning kit, and hearing protection you intend to use.")
paragraph(doc, "Familiarity matters more than novelty. Opening morning is not the time to learn a new safety, turret, bipod, sling, or magazine system. Practice moving into position, establishing a solid rest, ranging, settling the rifle, and staying on the rifle after the shot. If hearing protection is important to you, use electronic protection that reduces the rifle report while still allowing you to hear your guide. I cannot effectively spot a follow-up shot if you cannot hear my instructions.")
callout(doc, "Priority Safety Point", "As hunting guides, we will often lead from the front or hunt from the front while locating game, evaluating terrain, and positioning the hunter for a possible shot. Because a guide may be forward of the hunter, Wild Eyez Outfitters applies a safety policy that goes above and beyond state law, rules, and regulations: no live round is to be chambered into the barrel until the target animal has been positively identified, the decision to take the shot has been made, and the shooting lane to the front is entirely clear.")
callout(doc, "Ammunition Standard", "Lead ammunition is preferred. Copper ammunition is not preferred. I highly suggest .30 caliber or equivalent and a rifle system that you know well. It is your responsibility to understand your rifle's capabilities and ballistics. I will coach you through the actual shot, provide ranges, and operate as your spotter for follow-up shots.")
checklist(doc, "Firearms Checklist", [
    "Rifle you know well and are comfortable operating under pressure.",
    "Scope checked and rifle zeroed before arrival.",
    "Enough ammunition for the hunt and meaningful practice before arrival.",
    "Bipod or independent shooting sticks if part of your system.",
    "Cleaning kit and basic rifle tools.",
    "Electronic hearing protection if hearing protection is a concern.",
])

page_break(doc)
banner(doc, "PRESEASON PREPARATION")
paragraph(doc, "Your hunt begins before you travel to Utah. Wear your full clothing system, carry your loaded daypack, and practice realistic shooting positions. Pretend an animal has appeared and you have only seconds to move into an awkward uphill or downhill position. Get on the ground and look through your scope. These dry runs expose problems while there is time to correct them and create muscle memory for the moment when adrenaline makes clear thinking difficult.")
callout(doc, "Most Common Mistake", "The most common mistake is zeroing your weapon and practicing with it only at the range, punching paper on level ground. Wild Eyez recommends that you zero your centerfire rifle at 200 yards and/or your muzzleloader at 100 yards, then practice from realistic field positions.")
checklist(doc, "Preparation Checklist", [
    "Confirm your rifle is zeroed before arrival and verify the ammunition you plan to hunt with.",
    "Try your clothing and gear on at home, wear it around, and assume all potential firing positions.",
    "Wear your hunting clothing, orange outer-layer vest and headgear if required, daypack, accessory equipment, and rifle on walks or hikes before your hunt.",
    "Develop scenarios such as spotting a game animal, conducting a short stalk, and deploying your weapon system quickly from an abnormal or inconvenient field position.",
    "Practice shooting uphill and downhill. Nothing is level in the mountains.",
    "Practice retrieving specific items from your pack while wearing your hunting clothing.",
    "Tape or silence metal brackets, rings, buckles, trekking poles, and anything else that can clang during a stalk.",
])
paragraph(doc, "As a hunter, it is your responsibility, and as your Outfitter it is Wild Eyez' expectation, that you will perform this task at a minimum and be able to take game efficiently and decisively. We will respect and honor your game animal and will therefore expect from you your highest degree of ethical values and practice. All game animals are a beautiful treasure; Wild Eyez Outfitters sees it as exactly that.")
paragraph(doc, "A wounded animal can occur, but it is an ugly thing. It is not taken lightly by our guides, and it weighs heavily on one's mind. We will go to extremes, short of risking safety, to recover an animal that is wounded at your hands. Please do your best to minimize that outcome. Be prepared.")

page_break(doc)
banner(doc, "FINAL REMINDERS")
checklist(doc, "Priority Items", [
    "Bring your hunting license, permit, and required identification.",
    "Carry emergency contact information and any required prescription medications.",
    "Text Wild Eyez Outfitters first, when possible, since mountain service can be limited and I am often out of coverage locating your big game animal.",
    "Arrive physically prepared, mentally sharp, and with your clothing, gear, rifle, and pack working together.",
    "Use the checklists above before packing and again before departure.",
])
heading(doc, "Are You Ready For The Mountain???", 13)
paragraph(doc, "Thank you very much and again, congratulations on your upcoming hunt. Prepare yourself as best as possible and your likelihood of success will be high. We have prepared diligently for your hunt, are ready for your arrival, and look forward to making tracks with you on the mountain. This opportunity is a gift and an experience worth remembering for the rest of your life. We share our experience because hunting is our passion and because helping a client realize a personal goal is deeply rewarding. You may arrive as a stranger, but our hope is that you leave as a friend.")
p = paragraph(doc, "", before=4, after=4)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("PREPARE. HUNT SAFELY. HUNT ETHICALLY.")
set_font(r, size=14, bold=True, color=GREEN, name="Arial")
p = paragraph(doc, "", after=5)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("~SEE THE WILD!!!")
set_font(r, size=14, bold=True, italic=True, color=MAROON, name="Georgia")

card = ASSETS / "business-card.png"
if card.exists():
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t.columns[0].width = Inches(6.8)
    c = t.cell(0, 0)
    shade(c, "FFFFFF")
    cell_margins(c, 95, 95, 95, 95)
    cell_border(
        c,
        top={"val": "single", "sz": "8", "color": BRONZE},
        left={"val": "single", "sz": "8", "color": BRONZE},
        bottom={"val": "single", "sz": "8", "color": BRONZE},
        right={"val": "single", "sz": "8", "color": BRONZE},
    )
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(card), width=Inches(6.5))

for sec in doc.sections:
    p = sec.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Wild Eyez Outfitters | Manti, Utah | 435-851-6480 | tyler@wildeyez.net")
    set_font(r, size=8, color="555555", name="Arial")

doc.save(OUT_DOCX)
print(OUT_DOCX)
